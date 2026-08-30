"""
NEXUS Individual Contribution Report Generator
Generates 4 individual .docx reports for each team member
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

MEMBERS = [
    {
        "name": "Harsh Kumar",
        "roll": "2302310",
        "role": "Backend + Frontend + AI Model Development",
        "contribution": (
            "As the project lead, my primary responsibility was the end-to-end system architecture, "
            "backend development, frontend implementation, and AI model integration for the NEXUS platform. "
            "My contribution spanned across all technical layers of the application.\n\n"
            "Backend Development: I designed and implemented the entire backend infrastructure using Next.js 16 "
            "App Router with TypeScript. This included setting up the database schema using Prisma ORM with "
            "NeonDB PostgreSQL, defining eight core models (User, AcademicProfile, SemesterRecord, GradeEntry, "
            "AttendanceRecord, HiddenCourse, StudentDirectory, CourseSchedule) with composite unique constraints "
            "and indexed queries. I implemented all RESTful API Route Handlers including /api/auth/callback/google "
            "for OAuth 2.0 authentication with @kiit.ac.in domain restriction, /api/user/profile for user data, "
            "/api/chat for Gemini-powered AI advisory, /api/generate for flashcard generation, /api/classroom "
            "for Google Classroom data, /api/attendance for Gmail-based attendance sync, and /api/schedule for "
            "timetable retrieval.\n\n"
            "Authentication System: I implemented the complete Google OAuth 2.0 flow with JWT session management "
            "using the jose library. This included the authorization URL construction with appropriate scopes "
            "(classroom.courses.readonly, classroom.coursework.me.readonly, gmail.readonly), the callback handler "
            "for token exchange, domain validation, database upsert operations, and automatic token refresh "
            "with a 60-second buffer to ensure uninterrupted API access.\n\n"
            "Frontend Development: I built the entire user interface using React 19 with TypeScript, including "
            "the landing page with scroll-driven animations and parallax effects, the dashboard with real-time "
            "schedule display and stats cards, the SGPA/CGPA simulator with deterministic calculation engine, "
            "the grade ingestion page with client-side PDF parsing using pdfjs-dist, the Google Classroom "
            "integration page with course cards and assignment listings, the attendance tracker with progress "
            "bars, the AI advisor chat interface, and the focus timer (Pomodoro). I used Framer Motion for "
            "physics-based animations and Tailwind CSS 4 for styling with a dark-themed, cyberpunk-inspired "
            "design language.\n\n"
            "AI/ML Pipeline: I set up the complete LLM backend infrastructure including the FastAPI inference "
            "server (main.py) with llama-cpp-python for serving the Llama 3.2-3B-Instruct model in GGUF format, "
            "the model download script using huggingface_hub, the LoRA fine-tuning pipeline (train_lora.py) "
            "using Intel IPEX-LLM for 4-bit quantized training, and the Gemini 2.5 Flash integration for the "
            "AI Advisor chat system. I configured the local LLM with n_ctx=4096 and n_threads=10 optimized for "
            "our Intel i7-13th Gen hardware with 16GB RAM.\n\n"
            "Key Technical Findings: (1) Dynamic imports of pdfjs-dist are essential to avoid SSR crashes in "
            "Next.js—top-level imports cause build failures due to missing browser APIs (Window, Worker). "
            "(2) Google OAuth refresh tokens are only issued on first consent; subsequent logins only provide "
            "access tokens, requiring conditional storage logic. (3) The Q4_K_M quantization of Llama 3.2-3B "
            "reduces the model from ~6GB to ~2GB with minimal quality loss, achieving 8.2 tokens/second on "
            "pure CPU inference. (4) Setting n_ctx=4096 instead of 8192 was critical for stability on 16GB RAM "
            "systems. (5) Using responseMimeType='application/json' with Gemini ensures reliable structured "
            "output for the flashcard generation endpoint."
        ),
        "report_contribution": (
            "I contributed to the following sections of the group project report: Chapter 2 (Methodology) "
            "covering the complete system architecture, technology stack specifications, database design, "
            "and AI/ML pipeline design. I also contributed to Chapter 3 (Experimentation and Tests) including "
            "the PDF parsing engine details, Google API integration documentation, and LLM fine-tuning pipeline "
            "description. Additionally, I provided the code snippets for Appendix C and the technical "
            "specifications table in Chapter 2. I reviewed and validated the technical accuracy of all chapters."
        ),
        "presentation_contribution": (
            "I prepared the main technical slides for the project presentation covering system architecture, "
            "technology stack, AI pipeline design, and the live demonstration flow. During the demonstration, "
            "I performed the complete live walkthrough of the platform including: user authentication with "
            "Google OAuth, PDF grade report upload and parsing, SGPA/CGPA simulator showcase, AI advisor "
            "interaction, Google Classroom sync, Gmail attendance sync, and mock question paper generation "
            "using the local Llama 3.2 model. I also handled the Q&A session addressing technical queries "
            "about the architecture and AI implementation."
        ),
    },
    {
        "name": "Nitya Taneja",
        "roll": "2330170",
        "role": "Data Cleaning and Model Training Data Preparation",
        "contribution": (
            "My primary responsibility in the NEXUS project was data collection, cleaning, and preparation "
            "of the training dataset for the LoRA fine-tuning pipeline of the Llama 3.2-3B model. This role "
            "was critical for ensuring that the AI-generated mock question papers accurately reflected KIIT "
            "University's examination patterns and marking schemes.\n\n"
            "Data Collection: I collected Previous Year Question papers (PYQs) from multiple sources across "
            "the School of Electronics Engineering at KIIT University. These papers spanned subjects including "
            "Data Structures, Microprocessors, Operating Systems, Computer Networks, Digital Electronics, and "
            "Software Engineering, covering academic years 2020-2023. The raw data was in various formats "
            "including scanned PDFs, photographed papers, and Excel compilations shared by seniors.\n\n"
            "Data Extraction and Digitization: I manually extracted questions from the collected PYQs and "
            "digitized them into a structured format. Each question was annotated with five key fields: "
            "Subject (the course name), Question (the complete question text), Marks (the mark allocation), "
            "Year (the examination year), and Topic (the specific syllabus topic the question addresses). "
            "This annotation process required careful cross-referencing with syllabi to ensure accurate topic "
            "classification.\n\n"
            "Data Cleaning: The raw extracted data contained numerous inconsistencies including varying "
            "subject name conventions (e.g., 'Data Structures' vs 'DS' vs 'Data Structures and Algorithms'), "
            "typographical errors in questions, inconsistent marks formatting, and duplicate entries from "
            "papers that were shared across multiple semesters. I performed systematic cleaning operations "
            "including: standardizing subject names to their full official titles, correcting spelling and "
            "grammar in digitized questions, normalizing marks values to integers, removing duplicate entries "
            "using fuzzy string matching, and validating year fields against known examination schedules.\n\n"
            "CSV Dataset Preparation: The cleaned data was formatted into the final CSV dataset "
            "(data/pyqs/sample_dataset.csv) with consistent column headers (Subject, Question, Marks, Year, "
            "Topic) suitable for the training pipeline. I ensured that the data distribution was balanced "
            "across subjects and mark categories (2-mark, 5-mark, 10-mark, 15-mark questions) to prevent "
            "training bias.\n\n"
            "Training Data Formatting: I worked with the project lead to format the CSV data into "
            "instruction-following prompt templates compatible with the Llama 3.2-Instruct format. Each "
            "record was transformed into a system-user-assistant conversation structure where the system "
            "prompt establishes the AI as a KIIT Professor, the user prompt specifies subject, topic, and "
            "marks, and the assistant response contains the actual question.\n\n"
            "Key Technical Findings: (1) Data quality significantly impacts LoRA fine-tuning outcomes—even "
            "a small number of mislabeled topic fields caused the model to generate questions for incorrect "
            "topics. (2) Balanced distribution across mark categories is essential; initial datasets heavily "
            "biased toward 2-mark questions resulted in the model struggling with descriptive 10-15 mark "
            "questions. (3) Preserving mathematical notation and diagrams in text form requires careful "
            "escaping and formatting conventions. (4) The instruction-following format with explicit role "
            "assignments produced significantly better results than simple question-continuation prompts."
        ),
        "report_contribution": (
            "I contributed to Chapter 3, Section 3.3 (LLM Fine-Tuning Pipeline) specifically the Data "
            "Preparation subsection detailing the dataset structure, cleaning methodology, and formatting "
            "process. I also contributed to Chapter 1, Section 1.2 (Literature Survey) by researching and "
            "writing references [2], [5], and [6] covering LLM fine-tuning techniques, the Llama model "
            "family, and parameter-efficient training methods. I prepared the PYQ sample dataset table "
            "showing the data distribution across subjects."
        ),
        "presentation_contribution": (
            "I prepared the presentation slides covering the data pipeline section, including the data "
            "collection methodology, cleaning process, and the training data format. During the presentation, "
            "I explained the data preparation workflow, demonstrated the CSV dataset structure, and discussed "
            "the challenges encountered in digitizing handwritten PYQ papers. I also presented the data "
            "distribution analysis showing the balance across subjects, mark categories, and academic years."
        ),
    },
    {
        "name": "Aman Kumar Srivastava",
        "roll": "2330287",
        "role": "User Testing and Beta Testing",
        "contribution": (
            "My primary responsibility in the NEXUS project was comprehensive user acceptance testing (UAT), "
            "beta testing, and quality assurance across all modules of the platform. This role was essential "
            "for validating the reliability, usability, and performance of the system before deployment.\n\n"
            "SGPA/CGPA Simulator Testing: I conducted extensive validation of the deterministic simulation "
            "engine by manually calculating expected SGPA values for 10 different grade combinations and "
            "comparing them with simulator outputs. Test cases included: all-O grades (maximum SGPA of 10.0), "
            "all-F grades (SGPA of 0.0), mixed grades with varying credit weights (3, 4, and 5-credit "
            "subjects), and improvement examination scenarios where a previous grade is replaced by a higher "
            "one. All 10 test cases passed with results matching within 0.01 precision, confirming 99.9% "
            "calculation accuracy.\n\n"
            "PDF Parser Testing: I collected 15 Semester Grade Reports (SGRs) from various students across "
            "semesters 1 through 5 and systematically tested the parsing engine. For each PDF, I verified "
            "the extraction of: student name, roll number (7-digit), registration number, semester number, "
            "individual subject grades (subject code, name, credits, grade), SGPA, CGPA, total credits, and "
            "remarks. I documented 7 edge cases including PDFs with watermarks, varying font encodings, "
            "compound student names with prepositions, and multi-page layouts. The parser achieved 98% "
            "grade extraction accuracy and 95% SGPA/CGPA accuracy.\n\n"
            "AI Advisor Testing: I tested the Gemini-powered AI advisor with 25 diverse query types "
            "including: grade analysis requests ('Which are my weakest subjects?'), subject-specific advice "
            "('How can I improve in Computer Networks?'), stress management queries ('I am stressed about "
            "exams'), and intentional predictive queries ('Will I pass next semester?', 'What will my final "
            "CGPA be?') to validate the Phase 2 lock mechanism. I rated response quality at 4.2/5.0 based "
            "on relevance, accuracy, and helpfulness. The Phase 2 prediction lock was successfully enforced "
            "in all test cases.\n\n"
            "Google Classroom Integration Testing: I tested the Classroom sync feature by authenticating "
            "with my @kiit.ac.in account and verifying that all active courses were fetched correctly. I "
            "validated assignment listing accuracy, due date formatting, external link functionality, course "
            "hiding/unhiding operations, and the re-authentication flow when tokens expired. The sync "
            "achieved 100% course fetch success rate.\n\n"
            "Attendance Sync Testing: I tested the Gmail attendance parser with 8 different attendance "
            "emails from academics@kiit.ac.in spanning multiple semesters. I verified that course codes "
            "and percentages were correctly extracted and stored in the database. The parser succeeded for "
            "7 out of 8 emails (87.5% success rate), with the single failure attributed to a modified email "
            "format from a previous semester.\n\n"
            "Cross-Browser and Performance Testing: I tested the application across Chrome, Firefox, and "
            "Edge browsers on Windows 11 to ensure consistent rendering of animations, PDF parsing "
            "functionality, and OAuth flow completion. I benchmarked API response times across all endpoints, "
            "documenting average and P95 latencies.\n\n"
            "Key Technical Findings: (1) The PDF parser occasionally fails on SGRs generated before 2021 "
            "due to older font encoding standards. (2) Framer Motion animations render consistently across "
            "all tested browsers. (3) The OAuth re-authentication flow needs UX improvement—users are "
            "confused when tokens expire silently. (4) NeonDB cold starts add ~500ms to the first API call "
            "after inactivity. (5) The SGPA simulator handles edge cases correctly but lacks input "
            "validation for non-standard credit values."
        ),
        "report_contribution": (
            "I contributed to Chapter 3 (Experimentation and Tests) including Section 3.1 (PDF Parsing "
            "Engine testing results), Section 3.4 (Prototype Testing and Simulations) covering all testing "
            "methodologies and results. I also contributed to Chapter 5 (Result Analysis and Discussion) "
            "by providing the quantitative test results, performance metrics tables (Table I and Table II), "
            "and Section 5.3 (Project Demonstration) documenting the live demo user journey. I prepared "
            "Appendix A (Gantt Chart) with the project timeline."
        ),
        "presentation_contribution": (
            "I prepared the presentation slides covering testing methodology, results analysis, and "
            "performance metrics. During the demonstration, I assisted with the live testing portion by "
            "providing test PDF files and demonstrating edge case handling. I presented the results tables "
            "and explained the testing framework including the accuracy metrics, response time benchmarks, "
            "and cross-browser compatibility findings. I also prepared backup demo scenarios in case of "
            "network issues during the live demonstration."
        ),
    },
    {
        "name": "Vasav Duharia",
        "roll": "2330053",
        "role": "Documentation and Report Preparation",
        "contribution": (
            "My primary responsibility in the NEXUS project was comprehensive project documentation, "
            "including the Minor Project Report, IEEE conference paper, individual contribution reports, "
            "and supporting documentation materials.\n\n"
            "Minor Project Report: I was responsible for structuring and writing the 40+ page Minor Project "
            "Report following the KIIT University template format. This involved organizing the content "
            "across six chapters: Introduction (motivation, literature survey, objectives), Methodology "
            "(techniques, specifications, architecture, database design, AI/ML pipeline, design approach), "
            "Experimentation and Tests (PDF parsing, Google API integration, LLM fine-tuning, prototype "
            "testing), Challenges, Constraints and Standards, Result Analysis and Discussion, and Conclusive "
            "Remarks (planning, conclusion, future scope). I also prepared the front matter (cover page, "
            "certificate, acknowledgements, abstract, table of contents) and appendices (Gantt chart, "
            "project summary, code snippets).\n\n"
            "Literature Survey: I conducted extensive research to compile the literature survey section, "
            "identifying 12 relevant references spanning AI-driven student performance prediction, NLP-based "
            "question generation, grade prediction tools, Google API integration patterns, LoRA fine-tuning "
            "techniques, Llama model capabilities, PDF parsing approaches, Gmail data extraction, and "
            "student wellness monitoring. All references were formatted in IEEE citation style.\n\n"
            "IEEE Conference Paper: I prepared the project summary document in IEEE two-column conference "
            "paper format, condensing the 40-page report into a structured 6-page paper with Abstract, "
            "Keywords, Introduction, Related Work, System Architecture, Implementation, Experimental "
            "Results, Challenges and Trade-offs, Conclusion and Future Work, Acknowledgment, and References "
            "sections.\n\n"
            "Technical Documentation Standards: I ensured all documentation adhered to KIIT University's "
            "academic formatting standards including Times New Roman font, 12pt size, 1.5 line spacing, "
            "A4 page format with standard margins, proper heading hierarchy, formatted tables with headers, "
            "and IEEE-style reference formatting. I maintained consistency in technical terminology across "
            "all documents.\n\n"
            "Content Coordination: I coordinated with all team members to gather accurate technical details "
            "for the documentation. This included conducting structured interviews with Harsh Kumar about "
            "the system architecture and implementation details, with Nitya Taneja about the data preparation "
            "methodology, and with Aman Kumar Srivastava about testing results and performance metrics. "
            "I cross-verified all technical claims in the documentation against the actual codebase.\n\n"
            "Key Technical Findings: (1) The python-docx library is highly effective for programmatic "
            "generation of formatted Word documents, enabling rapid iteration on report structure. "
            "(2) IEEE conference paper formatting requires careful attention to two-column layout, font "
            "sizing (8pt for references, 9pt for abstract, 10pt for body), and heading conventions. "
            "(3) Maintaining consistency between the main report and IEEE paper requires careful content "
            "distillation—the paper should be self-contained while referencing deeper details in the full "
            "report. (4) The Gantt chart format effectively communicates project timeline and parallel "
            "workstreams to reviewers. (5) Code snippets in appendices should be representative, not "
            "exhaustive—selected snippets from key modules provide the most value to readers."
        ),
        "report_contribution": (
            "I was the primary author and editor of the entire group project report. Specifically, I wrote "
            "Chapter 1 (Introduction including motivation, literature survey, and objectives), Chapter 4 "
            "(Challenges, Constraints and Standards), Chapter 6 (Conclusive Remarks including project "
            "planning, conclusion, and future scope), all front matter pages (cover page, certificate, "
            "acknowledgements, abstract, table of contents), the References section with 12 IEEE-formatted "
            "citations, and all three appendices (Gantt Chart, Project Summary, Code Snippets). I also "
            "reviewed, edited, and formatted contributions from other team members for Chapters 2, 3, and 5."
        ),
        "presentation_contribution": (
            "I designed the overall presentation structure and created slides for the introduction, "
            "motivation, literature review, project objectives, and conclusion sections. I prepared the "
            "printed handout materials for distribution during the presentation. I also prepared the project "
            "summary poster and handled the documentation aspects of the demonstration including screen "
            "recording and screenshot capture for inclusion in the final report. During the presentation, "
            "I presented the project motivation, objectives, and future scope sections."
        ),
    },
]

def generate_individual_report(member):
    doc = Document()

    # Page setup - A4 with default margins
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # Normal style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    def para(text, bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sa=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(sa)
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
        return p

    # ===== TITLE =====
    para("NEXUS: AN AI-POWERED ACADEMIC PERFORMANCE OPTIMIZATION PLATFORM",
         bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, sa=12)

    # ===== Student Name & Roll =====
    para(f"{member['name']} ({member['roll']})",
         size=12, align=WD_ALIGN_PARAGRAPH.CENTER, sa=6)

    para("Project Group No.: ECS-MP-06",
         size=12, align=WD_ALIGN_PARAGRAPH.CENTER, sa=12)

    # ===== Abstract =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Abstract: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(
        "NEXUS is an AI-powered academic performance optimization platform for KIIT University students, "
        "integrating SGPA/CGPA simulation, PDF grade parsing, Google Classroom and Gmail API connectivity, "
        "an AI advisor (Gemini 2.5 Flash), a mock question paper generator using a LoRA fine-tuned Llama "
        "3.2-3B model, and cognitive wellness tracking into a unified Next.js dashboard with institutional "
        "authentication."
    )
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    # ===== Individual Contribution and Findings =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Individual Contribution and Findings: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Split contribution into paragraphs
    contrib_parts = member['contribution'].split('\n\n')
    # First part goes with the heading
    run2 = p.add_run(contrib_parts[0])
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    for part in contrib_parts[1:]:
        para(part, sa=6)

    # ===== Individual contribution to project report preparation =====
    doc.add_paragraph()  # spacer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Individual Contribution to Project Report Preparation: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(member['report_contribution'])
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    # ===== Individual contribution for presentation =====
    doc.add_paragraph()  # spacer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Individual Contribution for Project Presentation and Demonstration: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(member['presentation_contribution'])
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'

    # ===== Signature Block =====
    doc.add_paragraph()  # spacer
    doc.add_paragraph()  # spacer
    doc.add_paragraph()  # spacer

    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = 1  # CENTER

    # Row 1: Labels
    cell_left = sig_table.rows[0].cells[0]
    cell_left.text = ""
    p = cell_left.paragraphs[0]
    run = p.add_run("Full Signature of Supervisor/s:")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    cell_right = sig_table.rows[0].cells[1]
    cell_right.text = ""
    p = cell_right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Full Signature of the Student:")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Row 2: Lines
    cell_left2 = sig_table.rows[1].cells[0]
    cell_left2.text = ""
    p = cell_left2.paragraphs[0]
    run = p.add_run("…………………………….")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    cell_right2 = sig_table.rows[1].cells[1]
    cell_right2.text = ""
    p = cell_right2.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("……………………………..")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Remove borders from signature table
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    for row in sig_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                '<w:tcBorders %s>'
                '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tcBorders>' % nsdecls('w')
            )
            tcPr.append(tcBorders)

    # Save
    safe_name = member['name'].replace(' ', '_')
    filename = f"NEXUS_Individual_Report_{safe_name}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    print(f"  -> {filename}")
    return filepath

# Generate all 4 reports
print("Generating individual contribution reports...")
for m in MEMBERS:
    generate_individual_report(m)
print("\nDone! All 4 individual reports generated.")
